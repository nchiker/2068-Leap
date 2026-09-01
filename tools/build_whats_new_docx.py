#!/usr/bin/env python3
"""Build the styled Word edition of the Release 1 Beta What's New guide."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from build_user_manual_docx import configure, parse

ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / "docs" / "whats_new_release_1_beta.md"
OUTPUT = ROOT / "docs" / "2068-Leap_Whats_New_Release_1_Beta.docx"


def main():
    document = Document()
    configure(document)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("What's New")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("2068-Leap Release 1 Beta\nSince Public Preview 1")
    run.font.size = Pt(17)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    edition = document.add_paragraph()
    edition.alignment = WD_ALIGN_PARAGRAPH.CENTER
    edition.add_run(date.today().isoformat()).italic = True
    document.add_page_break()

    # Release notes read continuously; forcing every topic onto a new page
    # creates large white gaps and makes this short document feel fragmented.
    parse(document, SOURCE.read_text().splitlines(), page_break_sections=False)
    document.core_properties.title = "What's New in 2068-Leap Release 1 Beta"
    document.core_properties.subject = "Changes since Public Preview 1"
    document.core_properties.author = "2068-Leap Project"
    document.core_properties.keywords = (
        "2068-Leap, TS2068, Timex Sinclair, Release 1 Beta, What's New"
    )
    document.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
