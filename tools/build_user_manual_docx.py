#!/usr/bin/env python3
"""Build the styled Word edition of docs/user_manual.md."""

import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / "docs" / "user_manual.md"
OUTPUT = ROOT / "docs" / "2068_Leap_Users_Manual.docx"


def shade(cell_or_paragraph, fill):
    props = (
        cell_or_paragraph._tc.get_or_add_tcPr()
        if hasattr(cell_or_paragraph, "_tc")
        else cell_or_paragraph._p.get_or_add_pPr()
    )
    tag = "w:shd"
    node = props.find(qn(tag))
    if node is None:
        node = OxmlElement(tag)
        props.append(node)
    node.set(qn("w:fill"), fill)


def add_field(paragraph, instruction):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate.set(qn("w:dirty"), "true")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()._r
    run.extend((begin, text, separate, end))


def add_inline(paragraph, text):
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\[[^]]+\]\([^)]+\))")
    cursor = 0
    for match in pattern.finditer(text):
        paragraph.add_run(text[cursor : match.start()])
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.style = "Code Inline"
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            label, target = re.match(r"\[([^]]+)\]\(([^)]+)\)", token).groups()
            run = paragraph.add_run(label)
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            run.underline = True
            run.font.size = Pt(9)
            run._r.set(qn("w:rsidRPr"), target[:8].encode().hex()[:8])
        cursor = match.end()
    paragraph.add_run(text[cursor:])


def configure(document):
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.7)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in (
        ("Title", 30, RGBColor(0x1F, 0x4E, 0x79)),
        ("Heading 1", 20, RGBColor(0x1F, 0x4E, 0x79)),
        ("Heading 2", 15, RGBColor(0x2E, 0x75, 0xB6)),
        ("Heading 3", 12, RGBColor(0x3F, 0x3F, 0x3F)),
    ):
        style = document.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True

    code = document.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = "Consolas"
    code.font.size = Pt(8.5)
    code.paragraph_format.left_indent = Inches(0.2)
    code.paragraph_format.right_indent = Inches(0.1)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(7)

    inline = document.styles.add_style("Code Inline", WD_STYLE_TYPE.CHARACTER)
    inline.font.name = "Consolas"
    inline.font.size = Pt(9)
    inline.font.color.rgb = RGBColor(0x9C, 0x27, 0x2B)

    toc = document.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)
    toc.font.name = "Aptos"
    toc.font.size = Pt(10.5)
    toc.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    toc.paragraph_format.left_indent = Inches(0.15)
    toc.paragraph_format.space_after = Pt(2)

    for sec in document.sections:
        header = sec.header.paragraphs[0]
        header.text = "TIMEX SINCLAIR 2068  •  REDESIGNED BASIC ROM"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header.runs[0].font.size = Pt(8)
        header.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        footer = sec.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("TS2068 BASIC User's Manual   •   ")
        add_field(footer, "PAGE")


def parse(document, lines):
    index = 0
    in_code = False
    code_lines = []
    skip_contents = False
    first_h1 = True
    while index < len(lines):
        line = lines[index]
        if line.startswith("```"):
            if in_code:
                paragraph = document.add_paragraph(style="Code Block")
                paragraph.add_run("\n".join(code_lines))
                shade(paragraph, "F2F4F7")
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if line == "## Contents":
            skip_contents = True
            index += 1
            continue
        if skip_contents:
            if line == "---":
                skip_contents = False
            index += 1
            continue
        if not line or line == "---":
            index += 1
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            if level == 1 and first_h1:
                first_h1 = False
            else:
                if level == 2:
                    document.add_page_break()
                paragraph = document.add_heading(level=level)
                add_inline(paragraph, heading.group(2))
            index += 1
            continue

        if line.startswith("|") and index + 1 < len(lines) and re.match(
            r"^\|?\s*:?-+", lines[index + 1]
        ):
            rows = []
            while index < len(lines) and lines[index].startswith("|"):
                rows.append([part.strip() for part in lines[index].strip("|").split("|")])
                index += 1
            rows.pop(1)
            table = document.add_table(rows=len(rows), cols=max(map(len, rows)))
            table.style = "Light Shading Accent 1"
            for row_number, values in enumerate(rows):
                for column, value in enumerate(values):
                    cell = table.cell(row_number, column)
                    add_inline(cell.paragraphs[0], value)
                    if row_number == 0:
                        shade(cell, "1F4E79")
                        for run in cell.paragraphs[0].runs:
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.bold = True
            continue

        bullet = re.match(r"^[-*]\s+(.+)$", line)
        number = re.match(r"^\d+\.\s+(.+)$", line)
        if bullet or number:
            item = (bullet or number).group(1)
            index += 1
            while index < len(lines) and re.match(r"^\s{2,}\S", lines[index]):
                item += " " + lines[index].strip()
                index += 1
            paragraph = document.add_paragraph(
                style="List Bullet" if bullet else "List Number"
            )
            add_inline(paragraph, item)
            continue

        parts = [line.strip()]
        index += 1
        while index < len(lines):
            candidate = lines[index]
            if (
                not candidate
                or candidate.startswith(("#", "```", "|", "- ", "* "))
                or candidate == "---"
                or re.match(r"^\d+\.\s", candidate)
            ):
                break
            parts.append(candidate.strip())
            index += 1
        paragraph = document.add_paragraph()
        add_inline(paragraph, " ".join(parts))


def main():
    output = Path(sys.argv[1]) if len(sys.argv) > 1 else OUTPUT
    document = Document()
    configure(document)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("TS2068 BASIC")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("User's Manual\nRedesigned Timex Sinclair 2068 ROM")
    run.font.size = Pt(17)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    document.add_paragraph("\n")
    edition = document.add_paragraph()
    edition.alignment = WD_ALIGN_PARAGRAPH.CENTER
    edition.add_run(f"Project edition • {date.today().isoformat()}").italic = True
    document.add_page_break()

    document.add_heading("Contents", level=1)
    toc = document.add_paragraph()
    add_field(toc, 'TOC \\o "1-3" \\h \\z \\u')
    note = document.add_paragraph(
        "If page numbers are not visible, open the document in Word or "
        "LibreOffice and update the table of contents field."
    )
    note.runs[0].italic = True
    note.runs[0].font.size = Pt(9)
    for heading in re.findall(r"^##\s+(.+)$", SOURCE.read_text(), re.MULTILINE):
        if heading == "Contents":
            continue
        paragraph = document.add_paragraph(style="TOC 1")
        paragraph.add_run(heading)
    document.add_page_break()

    parse(document, SOURCE.read_text().splitlines())
    document.core_properties.title = "TS2068 BASIC User's Manual"
    document.core_properties.subject = "Redesigned Timex Sinclair 2068 ROM"
    document.core_properties.author = "TS2068 Redesigned ROM Project"
    document.core_properties.keywords = "TS2068, Timex Sinclair, BASIC, ROM"
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    print(f"Wrote {output}")


if __name__ == "__main__":
    main()
